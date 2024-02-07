import pandas as pd
import matplotlib.pyplot as plt
import os
import io
import docx
from docx.shared import Inches
from datetime import datetime
from decimal import Decimal, ROUND_DOWN
from pathlib import Path

# Dic criado para pegar os maiores valores da PV1 e PV2
dic_valor1 = {}
dic_valor2 = {}

# Lista criada para pegar tods os valroes de potência do dia.
potencia_total = []

# Dics criados para
lista1 = {}
lista2 = {}

# Lista para pegar os valores da área das PVs 1 e 2
gerado_pv1 = []
gerado_pv2 = []

# Dic criado para pegar a potência do dia como valor e o nome do arquivo como chave
gerado_dia = {}

# Cria uma lista de arquivos Excel no diretório especificado
caminho_diretorio = Path("C:\\Users\\Vertys\\Documents\\Export de dados")

# Lista para armazenar os caminhos dos arquivos
arquivos = []

# Itera sobre as pastas no diretório
for pasta in caminho_diretorio.iterdir():

    # Verifica se o item é um diretório
    if pasta.is_dir():

        # Itera sobre os arquivos dentro da pasta
        for arquivo in pasta.iterdir():
            # Adiciona o caminho do arquivo à lista
            arquivos.append(arquivo)

# Pegando cada excel da lista de arquivos
for arquivo in arquivos:

    try:

        # Ler Excel
        df = pd.read_excel(arquivo)

        # Potência do dia
        p = df['Yield(kWh)']
        pot = p.iloc[-10]

        # Pegando todos os dias de geração e sua geração e atribuindo o nome do arquivo com o seu valor
        gerado_dia[arquivo] = pot

    except:
        pass

# Filtro para pegar a area do melhor dia
valores = list(gerado_dia.values())

# Pegando a lista e ordedando os valores de forma decrecente
valor = sorted(valores, reverse=True)

# Número de itens por grupo
tamanho_grupo = 5

# Lista para armazenar os grupos
grupos = []

# Itera sobre a lista original em passos de tamanho_grupo
for i in range(0, len(valor), tamanho_grupo):
    grupo = valor[i:i + tamanho_grupo]
    grupos.append(grupo)

# Reorganiza o dic para que os valores se tornem chaves e as chaves os valores
dict_revisado = {valor: chave for chave, valor in gerado_dia.items()}

# Cria um dic  vazio para colocar os dias bons e ruins.
dict_refinado = {}

# Criando um dic, com apenas os dias bons para plotar na ordem correta no word
dict_bom = {}

# Intera sobre os valores da lsita que tem os grupos de 5 em 5
for i in grupos:
    # Pega apenas os primeiros  e os últimos valores da lista pois esta ordenada do maior para o menor
    v_index = len(i) - 1
    valor1 = i[0]
    valor2 = i[v_index]

    # Cria as chaves do novo dic, pegando os valores que se tornaram as chaves
    key1 = dict_revisado[valor1]
    key2 = dict_revisado[valor2]

    # Cria os valores do novo dic, pegando as chaves craidas acima
    dict_refinado[key1] = valor1
    dict_refinado[key2] = valor2

# Cria um documento do Word
doc = docx.Document()

# Intera agora só sobre os exceis
for arquivo in arquivos:

    try:
        print(f'arquivo x {arquivo}')

        # Ler Excel
        df = pd.read_excel(arquivo)

        # Ler pv1
        pt01 = df['PV1 Power(W)']
        pt1 = pt01

        # Ler pv2
        pt02 = df['PV2 Power(W)']
        pt2 = pt02

        # Potência do dia
        p = df['Yield(kWh)']
        pot = p.iloc[-10]

        # Add os valores dos dias na lista
        potencia_total.append((pot))

        # Time do dia
        tempes = df['Time']
        tempo = [str(t) for t in tempes]

        # Compara a maior potencia da PV1-----
        valor_maior_1 = pt1[0]

        i = 0
        for valor in pt1:

            i += 1

            # Verifica se o vlaor no indice atual e maior que o valor no indice seguinte
            if valor > valor_maior_1:
                valor_maior_1 = valor

                time_1 = tempo[i]

                dic_valor1[valor_maior_1] = time_1

                maior_valor_data_1 = {valor_maior_1: time_1}


            else:
                pass

        # Compara a maior potencia da PV2-----

        valor_maior_2 = pt2[0]

        i = 0
        for valor in pt2:

            i += 1

            # Verifica se o vlaor no indice atual e maior que o valor no indice seguinte
            if valor > valor_maior_2:
                valor_maior_2 = valor

                time_2 = tempo[i]

                dic_valor2[valor_maior_2] = time_2

                maior_valor_data_2 = (valor_maior_2, time_2)

            else:

                pass

        # Verificar qual é o maior valor entre as pvs e calcular sua diferença.

        if valor_maior_1 > valor_maior_2:

            difer = valor_maior_1 - valor_maior_2

        else:

            difer = valor_maior_2 - valor_maior_1

        i = 0
        t = 0

        # Cria lista para os valores da área de cada PV
        area1 = []
        area2 = []

        try:

            # Intera sobre os valores de tempo para pegar o tempo exato de cada ponto de leitura
            for te in tempo:
                i += 1

                # Corta a data para pegar só os horarios do indice sequencial
                tempo1 = te[11:]

                # Corta a data para pegar só os horarios do indice seguinte apartri do sequencial
                tempo001 = tempo[i]
                tempo01 = tempo001[11:]

                # Configura em uma variavel o formato que sera organizado os horarios cortados
                f = '%H:%M:%S'

                # Formatar valores para ver sua diferença de tempo
                dif = (datetime.strptime(tempo01, f) - datetime.strptime(tempo1, f)).total_seconds()

                # Converte a diferneça em valor de horario
                total_tem = (dif / 60) / 60

                # Trucando um valor de horario, para pegar o vlaro com extas 2 casa decimais com arredondamento para baixo
                t_truncado = Decimal(str(total_tem)).quantize(Decimal('0.00'), rounding=ROUND_DOWN)

                # Valor truncado vira string, ent aqui convertemos para um valor flutuante
                t_truncado_float = float(t_truncado)

                # Divide o valor de tempo por /2, para inserir no calculo de area de um, trapezio que exige o valor dividido por 2
                tem = t_truncado_float / 2

                # Calculando a área da PV1
                v = pt1[t]

                # Calculando area da pv1
                valor_base1 = pt1[i]
                are1 = (((v + valor_base1) * tem) / 1000)

                # Trunca os valor da area1
                n1_truncado = Decimal(str(are1)).quantize(Decimal('0.00'), rounding=ROUND_DOWN)

                # Valor truncado vira string, ent aqui convertemos para um valor flutuante
                n1_truncado_float = float(n1_truncado)

                # Add um dic com a data como chave e valor é a area1
                lista1[te] = n1_truncado_float

                # Add no dic a area da pv1
                area1.append((n1_truncado_float))

                # Calculando a área da PV2
                v = pt2[t]

                valor_base2 = pt2[i]
                are2 = (((v + valor_base2) * tem) / 1000)

                # Trunca os valrod a area2
                n2_truncado = Decimal(str(are2)).quantize(Decimal('0.00'), rounding=ROUND_DOWN)

                # Valor truncado vira string, ent aqui convertemos para um valor flutuante
                n2_truncado_float = float(n2_truncado)

                # Add um dic com a data como chave e valor é a area2
                lista2[te] = n2_truncado_float

                # Add no dic a area da pv1
                area2.append((n2_truncado_float))

                t += 1

        except Exception as e:
            print(e)
            pass

        print(arquivo)

        # Soma da aréa pv1
        area1_somada = 0

        for valor in area1:
            area1_somada += valor

        # Soma da aréa pv2
        area2_somada = 0

        for valor in area2:
            area2_somada += valor

        # Trunca os valores a area1
        n1_truncado = Decimal(str(area1_somada)).quantize(Decimal('0.00'), rounding=ROUND_DOWN)

        # Valor truncado vira string, ent aqui convertemos para um valor flutuante
        n1_truncado_float = float(n1_truncado)

        # Printa os valores da area do dia e o maior ponto de valro neste dia
        print(f'Areá do dia PV1 {n1_truncado_float}')

        print(f'Maior valor da PV1 ={maior_valor_data_1}')

        # Add valor da area da pv1 na lista para gerar a geração do mês
        gerado_pv1.append((n1_truncado_float))

        # Trunca os valores a area2
        n2_truncado = Decimal(str(area2_somada)).quantize(Decimal('0.00'), rounding=ROUND_DOWN)

        # Valor truncado vira string, ent aqui convertemos para um valor flutuante
        n2_truncado_float = float(n2_truncado)

        # Printa os valores da area do dia e o maior ponto de valro neste dia
        print(f'Areá do dia PV2 {n2_truncado_float}')

        print(f'Maior valor da PV2 ={maior_valor_data_2}')

        # Add valor da area da pv2 na lista
        gerado_pv2.append((n2_truncado_float))

        # Define o tamanho do grafico
        plt.figure(figsize=(35, 8))

        if arquivo in dict_refinado:

            # Cria um gráfico com as duas variáveis no eixo x e o tempo no eixo y
            plt.plot(tempo, pt1, label='PV1')
            plt.plot(tempo, pt2, label='PV2')

            # Adiciona um título e rótulos aos eixos
            plt.title('Potência de painéis solares')
            plt.xlabel('Tempo (h)')
            plt.ylabel('Potência (W)')

            # Adiciona uma anotação
            plt.annotate(f'PV1 ={maior_valor_data_1}', (time_1, valor_maior_1), xytext=(1, 0.8),
                         textcoords='axes fraction', size=10, arrowprops={'arrowstyle': '->', 'color': 'blue'})

            # Adiciona uma anotação
            plt.annotate(f'PV2 ={maior_valor_data_2}', (time_2, valor_maior_2), xytext=(1, 1),
                         textcoords='axes fraction', size=10, arrowprops={'arrowstyle': '->', 'color': 'orange'})

            # Pega a data correspondente a este dia:
            data = tempo[0]

            # Descreve a data:
            doc.add_paragraph(f'A data de registro é  {data}')

            # Add a geração do dia:
            doc.add_paragraph(f'Essa é a geração do dia: {pot}')

            # Add a aréa do grafico da pv1:
            doc.add_paragraph(f'Essa é a área do grafico da PV1 :{n1_truncado_float}')

            # Add a aréa do grafico da pv2:
            doc.add_paragraph(f'Essa é a área do grafico da PV2 :{n2_truncado_float}')

            # Descreve o nome do arquivo:
            doc.add_paragraph(f'O nome do arquivo é :{arquivo}')

            # Salva a imagem do gráfico
            plt.savefig("image.png")

            with open("image.png", "rb") as f:
                buf1 = f.read()

            # Cria um objeto de arquivo na memória a partir da imagem
            img1 = io.BytesIO(buf1)

            # Insere a imagem no documento
            doc.add_picture(img1, width=Inches(8), height=Inches(4))

            # Mostra o gráfico
            plt.show()
            plt.close()




        else:

            pass




    except:

        pass

# Soma e printa valor total do dia.
soma_dias = 0

for valor in potencia_total:
    soma_dias += valor

print(f'Geração dos dias {soma_dias}')

# Soma e printa valor total do mês.
mes1 = 0

for valor in gerado_pv1:
    mes1 += valor

mes2 = 0

for valor in gerado_pv2:
    mes2 += valor

print(f'Geração no mês pela PV1 {mes1}')

print(f'Geração no mês pela PV2 {mes2}'),

# Salva o documento
doc.save("Documento_das_Áreas.docx")





